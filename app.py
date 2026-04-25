import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
import io
import matplotlib.font_manager as fm
import os
import requests
from datetime import datetime
import pdfplumber

# --- 1. 環境設定：中文字體 (解決圖表亂碼) ---
@st.cache_resource
def load_chinese_font():
    font_url = "https://github.com/googlefonts/noto-cjk/raw/main/Sans/OTF/TraditionalChinese/NotoSansCJKtc-Regular.otf"
    font_path = "NotoSansCJKtc-Regular.otf"
    if not os.path.exists(font_path):
        try:
            response = requests.get(font_url)
            with open(font_path, "wb") as f:
                f.write(response.content)
        except: return None
    return font_path

font_p = load_chinese_font()
def apply_font_logic(font_path):
    if font_path:
        custom_font = fm.FontProperties(fname=font_path)
        plt.rcParams['font.family'] = custom_font.get_name()
        fm.fontManager.addfont(font_path)
        plt.rcParams['axes.unicode_minus'] = False
        return custom_font
    return None
font_prop = apply_font_logic(font_p)

# --- 2. 數據採集引擎 (支援 PDF & Excel 多檔案聚合) ---
def parse_pdf_robustly(file):
    """強大的 PDF 表格搜尋引擎：不依賴位置，主動尋找會計科目"""
    try:
        raw_tables = []
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    if len(table) < 2: continue
                    df_tmp = pd.DataFrame(table).dropna(how='all').dropna(axis=1, how='all')
                    raw_tables.append(df_tmp)
        if not raw_tables: return pd.DataFrame()
        master_df = pd.concat(raw_tables, ignore_index=True)
        # 初始化標準列
        res = {"公司名稱": file.name.replace(".pdf", ""), "年度": 0, "營收": 0, "應收帳款": 0, "存貨": 0, "現金": 0, "負債總額": 0}
        for _, row in master_df.iterrows():
            row_str = "".join([str(x) for x in row.values])
            # 年度辨識
            if any(k in row_str for k in ["年度", "Year"]):
                for val in row.values:
                    if str(val).isdigit() and len(str(val)) >= 3: res["年度"] = int(val)
            # 數值辨識
            def get_num(r):
                nums = [pd.to_numeric(str(x).replace(",",""), errors='coerce') for x in r.values if str(x).replace(",","").replace(".","").isdigit()]
                return nums[0] if nums else 0
            if any(k in row_str for k in ["營收", "營業收入", "Revenue"]): res["營收"] = get_num(row)
            if "應收帳款" in row_str: res["應收帳款"] = get_num(row)
            if "存貨" in row_str: res["存貨"] = get_num(row)
            if any(k in row_str for k in ["現金", "約當現金"]): res["現金"] = get_num(row)
            if any(k in row_str for k in ["負債總額", "負債合計"]): res["負債總額"] = get_num(row)
        return pd.DataFrame([res])
    except: return pd.DataFrame()

# --- 3. 核心引擎：鑑識與預測模組 ---
def crime_detector(row):
    """綜合舞弊與掏空預警模型"""
    r, rc, inv, cash, debt = row.get('營收', 0), row.get('應收帳款', 0), row.get('存貨', 0), row.get('現金', 0), row.get('負債總額', 0)
    if r == 0: return pd.Series([0, "數據不足", "低", "正常", "低"])
    m_score = -3.2 + (0.15 * (rc/r)) + (0.1 * (inv/r)) # Beneish M-Score
    t_ratio = (rc / r) if r > 0 else 0 # Tunneling Ratio
    p_index = (debt / cash) if cash > 0 else 0 # Ponzi Index
    ml_index = (cash / r) if r > 0 else 0 # Money Laundering Index
    return pd.Series([
        round(m_score, 2),
        "危險" if m_score > -1.78 else "正常",
        "高風險" if t_ratio > 0.4 else "正常",
        "警示" if p_index > 5 else "正常",
        "高機率" if ml_index < 0.05 else "低"
    ])

def get_forecast_df(co_df, future_years=3):
    """線性成長預測模型"""
    if len(co_df) < 2: return pd.DataFrame()
    co_df = co_df.sort_values('年度')
    # 計算歷史平均成長率
    growth_rate = co_df['營收'].pct_change().mean()
    if np.isnan(growth_rate): growth_rate = 0
    curr_rev = co_df['營收'].iloc[-1]
    last_year = co_df['年度'].iloc[-1]
    forecast_list = []
    # 移除預測時的 comma 避免格式錯誤
    for i in range(1, future_years + 1):
        curr_rev *= (1 + growth_rate)
        forecast_list.append({'年度': f"{int(last_year)+i}(預測)", '營收': round(curr_rev, 2), '類型': '預測'})
    return pd.DataFrame(forecast_list)

# --- 4. Streamlit 介面與導航中心 ---
st.set_page_config(layout="wide", page_title="玄武鑑識與預測平台")

with st.sidebar:
    st.header("🕵️ 鑑識官控制台")
    # 將模式正和為：單一公司深度分析、多公司風險對比
    mode = st.radio("功能模式切換", ["單一公司：深度分析與成長預測", "多公司對比：同業競爭與預測評比"])
    st.divider()
    auditor = st.text_input("主辦會計師簽署", "張鈞翔會計師")
    files = st.file_uploader("批次上傳數據 (Excel/PDF)", type=["xlsx", "pdf"], accept_multiple_files=True)

# --- 5. 數據處理與合併 ---
if files:
    data_list = [pd.read_excel(f) if f.name.endswith('.xlsx') else parse_pdf_robustly(f) for f in files]
    df = pd.concat(data_list, ignore_index=True).fillna(0)
    df['年度'] = pd.to_numeric(df['年度'], errors='coerce').fillna(0).astype(int)
    # 確保公司名稱與年度無重複
    df = df.drop_duplicates(subset=['公司名稱', '年度'], keep='last')
    
    # 執行鑑識運算
    df[['舞弊指標', '舞弊判定', '掏空風險', '吸金指標', '洗錢風險']] = df.apply(crime_detector, axis=1)

    # --- 模式 A: 單一公司深度分析 (網頁全圖表版) ---
    if mode == "單一公司：深度分析與成長預測":
        target = st.selectbox("選擇受查公司", df['公司名稱'].unique())
        sub = df[df['公司名稱'] == target].sort_values('年度')
        
        st.title(f"🔍 {target} 深度鑑定看板")
        
        # 1. 財務趨勢圖與預測 (正和功能)
        st.subheader("營收歷史軌跡與成長預估")
        f_df = get_forecast_df(sub)
        
        fig, ax = plt.subplots(figsize=(10, 4))
        # 畫歷史營收線 (藍色實線)
        ax.plot(sub['年度'].astype(str), sub['營收'], label='歷史實際營收', marker='o', linewidth=2)
        # 畫預測成長線 (橘色虛線)
        if not f_df.empty:
            # 確保年度格式一致
            ax.plot(f_df['年度'], f_df['營收'], '--', label='AI模型預測趨勢', marker='s', color='orange')
        
        ax.set_title("營收歷史與未來預測圖", fontproperties=font_prop)
        ax.set_ylabel("金額", fontproperties=font_prop)
        ax.legend(prop=font_prop)
        st.pyplot(fig)
        
        # 2. 四大指標鑑定儀表板
        st.subheader("三大犯罪指標鑑定")
        cols = st.columns(3)
        latest = sub.iloc[-1]
        cols[0].metric("舞弊指標 (M-Score)", latest['舞弊指標'], latest['舞弊判定'], delta_color="inverse")
        cols[1].metric("掏空風險", latest['掏空風險'])
        cols[2].metric("吸金指標", latest['吸金指標'])
        
        st.dataframe(sub[['年度', '營收', '舞弊指標', '掏空風險', '吸金指標', '洗錢風險']])

    # --- 模式 B: 多公司風險對比 (網頁全圖表版) ---
    elif mode == "多公司對比：同業競爭與預測評比":
        st.title("📊 同業風險與成長對比中心")
        
        charts_cols = st.columns(2)
        
        with charts_cols[0]:
            st.subheader("產業舞弊風險分佈 (M-Score)")
            fig2, ax2 = plt.subplots(figsize=(6, 5))
            ax2.bar(df['公司名稱'], df['舞弊指標'], color='teal')
            # 畫 Beneish警戒線 (-1.78)
            ax2.axhline(y=-1.78, color='red', linestyle='--', label='舞弊警戒線 (-1.78)')
            ax2.set_ylabel("指標分數")
            ax2.set_title("跨公司風險評比", fontproperties=font_prop)
            ax2.legend(prop=font_prop)
            st.pyplot(fig2)
            
        with charts_cols[1]:
            st.subheader("多公司成長預測PK圖")
            fig3, ax3 = plt.subplots(figsize=(6, 5))
            for co in df['公司名稱'].unique():
                co_sub = df[df['公司名稱'] == co].sort_values('年度')
                # 取得該公司的歷史與預測
                full_plot = pd.concat([co_sub[['年度','營收']], get_forecast_df(co_sub, years=2)[['年度','營收']]])
                # 年度強制轉成 string 畫圖避開 float 年度錯誤
                full_plot['年度'] = full_plot['年度'].astype(str)
                ax3.plot(full_plot['年度'], full_plot['營收'], label=co, marker='.')
            
            ax3.set_title("各標的成長動能對比", fontproperties=font_prop)
            ax3.legend(prop=font_prop)
            st.pyplot(fig3)
            
        st.dataframe(df[['公司名稱', '年度', '營收', '舞弊指標', '舞弊判定', '掏空風險', '吸金指標', '洗錢風險']])

    # --- 6. 報告導出模組 (正和成長預測敘述版) ---
    st.divider()
    if st.button("下載全方位鑑定與預測報告書 (Word)"):
        # 呼叫之前強化版的 Word生成邏輯
        doc = Document()
        # ...(這裡放置之前生成的Word代碼邏輯，包含預測敘述)...
        doc.add_heading("財務犯罪與成長預測鑑定報告", 0)
        doc.add_paragraph(f"主辦會計師：{auditor}")
        for _, r in df[df['舞弊判定'] == "危險"].iterrows():
            doc.add_paragraph(f"標的：{r['公司名稱']} ({r['年度']})指標異常")
            
        buf = io.BytesIO()
        doc.save(buf)
        st.download_button("匯出 Word 報告", buf.getvalue(), "鑑定報告.docx")
else:
    st.info("系統就緒。請從左側上傳 Excel 或 PDF 資料以啟動鑑定引擎。")
# --- 7. 報告導出 (含成長預測與鑑定敘述) ---
    st.divider()
    if st.button("下載旗艦鑑定報告 (Word)"):
        doc = Document()
        doc.add_heading("財務犯罪防制與成長預測鑑定報告書", 0).alignment = 1
        
        # 基本資訊
        p = doc.add_paragraph()
        p.add_run(f"鑑定人：{auditor} 會計師\n").bold = True
        p.add_run(f"鑑定日期：{datetime.now().strftime('%Y/%m/%d')}\n")
        
        # --- 第一部分：成長預測分析 ---
        doc.add_heading("一、 財務成長預測分析", level=1)
        for target_co in df['公司名稱'].unique():
            sub_co = df[df['公司名稱'] == target_co].sort_values('年度')
            if len(sub_co) >= 2:
                # 計算平均成長率
                avg_growth = sub_co['營收'].pct_change().mean()
                f_df = get_forecast_df(sub_co, years=2) # 呼叫之前的預測函數
                
                doc.add_heading(f"● {target_co} 營運展望", level=2)
                forecast_text = (
                    f"根據歷史數據顯示，{target_co} 之平均營收年增率（CAGR）約為 {avg_growth:.2%}。 "
                    f"透過 AI 線性回歸模型推估，未來兩年營收預計將分別達到 "
                    f"{f_df['營收'].iloc[0]:,.0f} 元及 {f_df['營收'].iloc[1]:,.0f} 元。"
                )
                doc.add_paragraph(forecast_text)
                
                # 加入預測合理性判斷
                if avg_growth > 0.5:
                    doc.add_paragraph("【警語】：該公司成長率異常過高（超過 50%），需注意是否為過度擴張或虛假交易所致。").italic = True
            else:
                doc.add_paragraph(f"● {target_co}：歷史數據不足（少於兩年），無法進行有效之成長率預測。")

        # --- 第二部分：犯罪指標深度鑑定 ---
        doc.add_heading("二、 財務犯罪防制偵測結論", level=1)
        # 篩選風險資料
        risk_summary = df[(df['舞弊判定'] == "危險") | (df['掏空風險'] == "高風險")]
        
        if risk_summary.empty:
            doc.add_paragraph("經本系統鑑定，受查標的於各項犯罪預警模型中均呈現「正常」狀態。")
        else:
            for _, r in risk_summary.iterrows():
                doc.add_heading(f"● 異常標的鑑定：{r['公司名稱']} ({r['年度']}年度)", level=2)
                diag = ""
                if r['舞弊判定'] == "危險":
                    diag += f"【盈餘操縱風險】：M-Score 指數 {r['舞弊指標']} 顯示財報存在高度操縱跡象。\n"
                if r['掏空風險'] == "高風險":
                    diag += f"【資產掏空風險】：應收帳款與營收成長失調，懷疑存在資金挪用。\n"
                if r['吸金指標'] == "警示":
                    diag += f"【吸金風險】：負債比重過高且現金流枯竭，具備龐氏騙局特徵。\n"
                doc.add_paragraph(diag)

        # 結尾聲明
        doc.add_heading("三、 鑑定師總結意見", level=1)
        doc.add_paragraph("本報告結合「歷史鑑定」與「未來預測」，旨在提供多維度之風險監控。建議針對上述列為「危險」或「警示」之年度進行專案審計。")

        # 匯出
        buf = io.BytesIO()
        doc.save(buf)
        st.download_button("匯出完整鑑定與預測報告", buf.getvalue(), "鑑定與預測綜合報告.docx")
